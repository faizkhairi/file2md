"""Shared test fixtures — programmatic PDF/DOCX generation."""

from __future__ import annotations

from pathlib import Path

import pymupdf
import pytest
from docx import Document


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Generate a simple multi-page test PDF."""
    doc = pymupdf.open()

    # Page 1
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Company Report Header", fontsize=12)
    page1.insert_text((72, 110), "This is the first paragraph of the document.")
    page1.insert_text((72, 130), "It continues on the next line with more text.")
    page1.insert_text((72, 170), "Another paragraph follows after a gap.")
    page1.insert_text((72, 750), "Page 1 of 3")

    # Page 2
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Company Report Header", fontsize=12)
    page2.insert_text((72, 110), "Second page content begins here.")
    page2.insert_text((72, 130), "This page has a hyphen-")
    page2.insert_text((72, 150), "ated word split across lines.")
    page2.insert_text((72, 750), "Page 2 of 3")

    # Page 3
    page3 = doc.new_page()
    page3.insert_text((72, 72), "Company Report Header", fontsize=12)
    page3.insert_text((72, 110), "Final page of the document.")
    page3.insert_text((72, 750), "Page 3 of 3")

    path = tmp_path / "sample.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def single_page_pdf(tmp_path: Path) -> Path:
    """Generate a single-page PDF."""
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello World", fontsize=14)
    page.insert_text((72, 110), "This is a simple single-page PDF.")
    path = tmp_path / "single.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def scanned_pdf(tmp_path: Path) -> Path:
    """Generate a PDF with only an image (simulating scanned document)."""
    doc = pymupdf.open()
    page = doc.new_page()
    # Create a large image-like pixmap covering most of the page
    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 200, 200), 0)
    pix.set_rect(pix.irect, (200, 200, 200))
    page.insert_image(page.rect, pixmap=pix)
    path = tmp_path / "scanned.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Generate a test DOCX with headings, lists, and a table."""
    doc = Document()

    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is the first paragraph with some content.")

    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Another paragraph here.")

    # Bullet list
    doc.add_paragraph("First bullet item", style="List Bullet")
    doc.add_paragraph("Second bullet item", style="List Bullet")

    # Numbered list
    doc.add_paragraph("First numbered item", style="List Number")
    doc.add_paragraph("Second numbered item", style="List Number")

    doc.add_heading("Data Table", level=2)

    # Table
    table = doc.add_table(rows=3, cols=3)
    headers = ["Name", "Age", "City"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    data = [["Alice", "30", "KL"], ["Bob", "25", "Penang"]]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("This concludes the test document.")

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def formatted_docx(tmp_path: Path) -> Path:
    """Generate a DOCX with bold, italic, and mixed formatting."""
    doc = Document()
    doc.add_heading("Formatted Document", level=1)

    para = doc.add_paragraph()
    para.add_run("Normal text, ")
    run_bold = para.add_run("bold text")
    run_bold.bold = True
    para.add_run(", and ")
    run_italic = para.add_run("italic text")
    run_italic.italic = True
    para.add_run(".")

    path = tmp_path / "formatted.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def empty_pdf(tmp_path: Path) -> Path:
    """Generate a PDF with no text content."""
    doc = pymupdf.open()
    doc.new_page()
    path = tmp_path / "empty.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def pdf_with_4_line_headers(tmp_path: Path) -> Path:
    """Generate a PDF with 4-line headers that exceed n_check=3."""
    doc = pymupdf.open()

    for page_num in range(1, 5):
        page = doc.new_page()
        # 4-line header block
        page.insert_text((72, 50), "MyDigital Integration Guideline", fontsize=10)
        page.insert_text((72, 65), "SSO Service Provider", fontsize=10)
        page.insert_text((72, 80), "Document ID: PP24176 v5.0A", fontsize=8)
        page.insert_text((72, 95), "Confidential", fontsize=8)
        # Enough content lines for n_check to reach 4+ (need 20+ lines)
        y = 140
        for line_num in range(1, 18):
            page.insert_text((72, y), f"Content line {line_num} for page {page_num}.")
            y += 18
        # Footer
        page.insert_text((72, 750), f"Page {page_num} of 4")

    path = tmp_path / "headers_4_line.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def pdf_with_leading_spaces(tmp_path: Path) -> Path:
    """Generate a PDF with indented text (simulating enterprise layout)."""
    doc = pymupdf.open()
    page = doc.new_page()
    # Indented text at various X positions
    page.insert_text((150, 72), "3.1.1 Section Title", fontsize=12)
    page.insert_text((150, 100), "This paragraph is indented in the PDF layout.")
    page.insert_text((200, 130), "Even more indented sub-content.")
    page.insert_text((72, 170), "Normal content at standard position.")
    path = tmp_path / "leading_spaces.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def pdf_with_images(tmp_path: Path) -> Path:
    """Generate a PDF with text and a large image (mixed page)."""
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Figure 1: Architecture Diagram", fontsize=12)
    # Insert a large image covering ~30% of the page
    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 300, 300), 0)
    pix.set_rect(pix.irect, (100, 150, 200))
    page.insert_image(pymupdf.Rect(72, 100, 400, 400), pixmap=pix)
    page.insert_text((72, 420), "The diagram above shows the system architecture.")
    path = tmp_path / "with_images.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def pdf_with_toc(tmp_path: Path) -> Path:
    """Generate a PDF with TOC-style dot leader lines."""
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Table of Contents", fontsize=14)
    page.insert_text((72, 110), "Introduction .............. 5")
    page.insert_text((72, 130), "3.1 API Flow ......... 12")
    page.insert_text((72, 150), "4.0 Security ................. 28")
    page.insert_text((72, 180), "Regular content without dots.")
    path = tmp_path / "with_toc.pdf"
    doc.save(str(path))
    doc.close()
    return path
