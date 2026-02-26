"""DOCX to Markdown conversion using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .normalize import apply_normalization
from .utils import ConversionOptions, ConversionResult, ExitCode, build_metadata


def convert_docx(path: Path, options: ConversionOptions) -> ConversionResult:
    """Convert a DOCX file to Markdown.

    Args:
        path: Path to the DOCX file.
        options: Conversion options.

    Returns:
        ConversionResult with the markdown content.
    """
    try:
        doc = Document(str(path))
    except Exception as e:  # noqa: BLE001
        return ConversionResult(
            markdown="",
            source_file=path.name,
            warnings=[f"Failed to open DOCX: {e}"],
            exit_code=ExitCode.EXTRACTION_FAILED,
        )

    parts: list[str] = []

    # Iterate over document body elements in order
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is not None:
                md = _convert_paragraph(para)
                if md is not None:
                    parts.append(md)

        elif tag == "tbl":
            table = _find_table(doc, element)
            if table is not None:
                md = _convert_table(table)
                if md:
                    parts.append(md)

    markdown = "\n\n".join(parts)

    if options.clean:
        markdown = apply_normalization(markdown)

    # Add metadata header
    metadata = build_metadata(path, options)
    markdown = metadata + "\n\n" + markdown

    # Truncation
    if options.max_chars and len(markdown) > options.max_chars:
        markdown = (
            markdown[: options.max_chars]
            + f"\n\n[... truncated at {options.max_chars} characters]"
        )

    return ConversionResult(
        markdown=markdown,
        source_file=path.name,
        exit_code=ExitCode.SUCCESS,
    )


def _find_paragraph(doc: Any, element: object) -> Paragraph | None:
    """Find the Paragraph object matching an XML element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para  # type: ignore[no-any-return]
    return None


def _find_table(doc: Any, element: object) -> Table | None:
    """Find the Table object matching an XML element."""
    for table in doc.tables:
        if table._element is element:
            return table  # type: ignore[no-any-return]
    return None


def _convert_paragraph(para: Paragraph) -> str | None:
    """Convert a DOCX paragraph to Markdown."""
    style_name = (para.style.name or "").lower() if para.style else ""
    text = _extract_runs(para)

    if not text.strip():
        return None

    # Headings
    if style_name.startswith("heading"):
        level_str = style_name.replace("heading", "").strip()
        try:
            level = min(int(level_str), 6)
        except ValueError:
            level = 1
        prefix = "#" * level
        return f"{prefix} {text}"

    # List items — style-name based detection
    if _is_bullet_list(para, style_name):
        return f"- {text}"

    if _is_numbered_list(para, style_name):
        return f"1. {text}"

    # Normal paragraph
    return text


def _is_bullet_list(para: Paragraph, style_name: str) -> bool:
    """Detect if a paragraph is a bulleted list item."""
    # Style-name detection (python-docx created docs)
    if style_name.startswith("list bullet"):
        return True

    # XML numPr fallback (Word-created docs)
    p_pr = para._element.pPr
    if p_pr is not None:
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            num_id_el = num_pr.find(qn("w:numId"))
            if num_id_el is not None:
                # Check numFmt if possible, otherwise fall back to style hint
                fmt = _resolve_num_format(para, num_pr)
                if fmt == "bullet":
                    return True
    return False


def _is_numbered_list(para: Paragraph, style_name: str) -> bool:
    """Detect if a paragraph is a numbered list item."""
    if style_name.startswith("list number"):
        return True

    p_pr = para._element.pPr
    if p_pr is not None:
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            num_id_el = num_pr.find(qn("w:numId"))
            if num_id_el is not None:
                fmt = _resolve_num_format(para, num_pr)
                if fmt and fmt != "bullet":
                    return True
    return False


def _resolve_num_format(para: Paragraph, num_pr: object) -> str | None:
    """Try to resolve the numbering format from the document's numbering definitions.

    Returns 'bullet', 'decimal', etc., or None if resolution fails.
    """
    try:
        ilvl_el = num_pr.find(qn("w:ilvl"))  # type: ignore[attr-defined]
        num_id_el = num_pr.find(qn("w:numId"))  # type: ignore[attr-defined]
        if ilvl_el is None or num_id_el is None:
            return None

        ilvl = ilvl_el.get(qn("w:val"))
        num_id = num_id_el.get(qn("w:val"))
        if not ilvl or not num_id:
            return None

        # Access the numbering part
        numbering_part = para.part.numbering_part  # type: ignore[attr-defined]
        if numbering_part is None:
            return None

        numbering_elem = numbering_part._element

        # Find abstractNumId for this numId
        abstract_num_id = None
        for num in numbering_elem.findall(qn("w:num")):
            if num.get(qn("w:numId")) == num_id:
                abs_ref = num.find(qn("w:abstractNumId"))
                if abs_ref is not None:
                    abstract_num_id = abs_ref.get(qn("w:val"))
                break

        if abstract_num_id is None:
            return None

        # Find the numFmt for this abstractNum at this ilvl
        for abstract_num in numbering_elem.findall(qn("w:abstractNum")):
            if abstract_num.get(qn("w:abstractNumId")) == abstract_num_id:
                for lvl in abstract_num.findall(qn("w:lvl")):
                    if lvl.get(qn("w:ilvl")) == ilvl:
                        num_fmt = lvl.find(qn("w:numFmt"))
                        if num_fmt is not None:
                            return num_fmt.get(qn("w:val"))  # type: ignore[no-any-return]
                break
    except Exception:  # noqa: BLE001
        pass

    return None


def _extract_runs(para: Paragraph) -> str:
    """Extract text from paragraph runs with inline formatting."""
    parts: list[str] = []

    for run in para.runs:
        text = run.text
        if not text:
            continue

        # Apply inline formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        # Monospace detection
        if run.font.name and any(
            mono in run.font.name.lower()
            for mono in ("mono", "courier", "consolas", "menlo")
        ):
            text = f"`{text}`"

        parts.append(text)

    # Also check for hyperlinks in the paragraph XML
    result = "".join(parts)

    # If no runs produced text, fall back to paragraph text
    if not result.strip() and para.text.strip():
        return para.text

    return result


def _convert_table(table: Table) -> str:
    """Convert a DOCX table to GFM markdown table."""
    if not table.rows:
        return ""

    rows: list[str] = []

    for row in table.rows:
        cells: list[str] = []
        prev_element = None
        for cell in row.cells:
            # Skip horizontally merged cells within the same row
            if cell._element is prev_element:
                continue
            prev_element = cell._element
            cell_text = cell.text.strip().replace("|", "\\|").replace("\n", " ")
            cells.append(cell_text)
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) >= 1:
        # Insert separator after first row (header)
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)

    return "\n".join(rows)
